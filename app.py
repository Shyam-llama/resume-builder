import streamlit as st
from pdfminer.high_level import extract_text
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import json
import aiohttp
import asyncio
import time
import openai
import os
import re
import pandas as pd


api_key = st.secrets.api_key
API_KEY = api_key
openai.api_key = api_key


def convert_files_to_text(file_path):
    try:
        if file_path.endswith('.pdf'):
            text = convert_pdf_to_text2(file_path)
        elif file_path.endswith('.docx'):
            text = convert_docx_to_text(file_path)
        elif file_path.endswith('.txt'):
            text = convert_txt_to_text(file_path)
        else:
            return "Not a Resume"
        return text
    except Exception as e:
        print(f"Error converting file {file_path} to text: {e}")
        st.error(f"Error converting file {file_path} to text: {e}")
        return ""
def convert_pdf_to_text2(file_path):
    try:
        text = extract_text(file_path)
        return text
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {e}")
        return ""

def convert_docx_to_text(file_path):
    try:
        doc = docx.Document(file_path)
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\\n'
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
        return ""
def convert_txt_to_text(file_path):
    try:
        with open(file_path, 'r') as file:
            content = file.read()
        return content
    except Exception as e:
        return f"An error occurred: {e}"



def truncate_text_by_words(text, max_words=4000):
    """
    Truncates the text to a specified number of words.
    """
    words = text.split()
    if len(words) <= max_words:
        return text
    return " ".join(words[:max_words])


system1='''You are an excellent NLP engineer, skilled talent recruiter and data scientist and your task is to analyse and parse candidate resumes into meaningful structured JSON format.
        You will be provided with candidate resume text.
        The system instruction is:
        Step-1:
        Analyse and parse the following information from the candidate's resume, do not just extract the data, rephrase it meaningfully;
        return the meaningful parsed data in a sturctured JSON format with key and corresponding value format as follows-
        'name': string,
        'gmail': string,
        'phone number' : string,
        'social media links': list of string,
        If value of a key is missing in the resume then value should be null.
        If not a resume then all the key's value should be null.
        Step-2:
        Only return the parsed JSON format resume, nothing else.'''

system2='''You are an excellent NLP engineer, skilled talent recruiter and data scientist and your task is to analyse and parse candidate resumes into meaningful structured JSON format.
        You will be provided with candidate resume text.
        The system instruction is:
        Step-1:
        Analyse and parse the following information from the candidate's resume, do not just extract the data, rephrase it meaningfully;
        return the meaningful parsed data in a sturctured JSON format with key and corresponding value format as follows-
        'skillset and expertise': list of string,
        'certifications': list of string,
        'Explanation of projects': list of string under 200 tokens,
        'Explanation of position of responsibilities': list of string under 200 tokens,
        If value of a key is missing in the resume then value should be null.
        If not a resume then all the key's value should be null.
        Step-2:
        Only return the parsed JSON format resume, nothing else.'''

system3='''You are an excellent NLP engineer, skilled talent recruiter and data scientist and your task is to analyse and parse candidate resumes into meaningful structured JSON format.
        You will be provided with candidate resume text.
        The system instruction is:
        Step-1:
        Analyse and parse the following information from the candidate's resume, do not just extract the data, rephrase it meaningfully;
        return the meaningful parsed data in a sturctured JSON format with key and corresponding value format as follows-
        'years of experience': string,
        'Previous work experience description': list of string under 200 tokens,
        'educational qualification': list of string,
        'extracurriculars': list of string,
        'awards and achievements': list of string,
        'previous job title': list of string
        If value of a key is missing in the resume then value should be null.
        If not a resume then all the key's value should be null.
        Step-2:
        Only return the parsed JSON format resume, nothing else. '''

systems = [system1, system2, system3]


async def async_openai_request(session, resumetext, system):
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    prompt = f"""
            Only return the structured parsed json format of the resume of candidate.
            Information about the candidate's resume is given inside text delimited by triple backticks.

            Candidate's Resume :```{resumetext}```

            """
    data = {
        "model": "gpt-3.5-turbo-16k",
        "messages": [{"role": "system", "content": system}, {"role": "user", "content": prompt}],
        "temperature": 0
    }
    async with session.post(url, json=data, headers=headers) as response:
        return await response.json()


# Process responses from OpenAI API
def process_responses(responses):
    output_list = [json.loads(resp['choices'][0]['message']['content']) for resp in responses]   
    combined_dict = {k: v for response in output_list for k, v in response.items()}
    return combined_dict

# Main async function to fetch and process responses
async def fetch_and_process(resumetext, systems):
    async with aiohttp.ClientSession() as session:
        results = await asyncio.gather(*(async_openai_request(session, resumetext, system) for system in systems))
    return process_responses(results)



# Function to set space after a paragraph to zero
def set_space_after(paragraph, space):
    p_spacing = OxmlElement('w:spacing')
    p_spacing.set(qn('w:after'), str(space))
    paragraph._element.get_or_add_pPr().append(p_spacing)

# Function to set cell background color
def set_cell_background(cell, fill):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Function for creating a document from a JSON-like structure
def create_doc_from_json1(json_data, filename):
    doc = Document()
    
    # Set the page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Add a table for the header
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(4.25)
    header_table.columns[1].width = Inches(2.25)
    
    # Populate header table
    name_cell = header_table.cell(0, 0) 
    name_cell.text = f"{json_data['name']}\n{json_data['social media links'][0]}"
    set_cell_background(name_cell, 'ADD8E6')  # Light gray background
    contact_cell = header_table.cell(0, 1)
    contact_info = f"Contact No: {json_data['phone number']}\nEmail: {json_data['gmail']}\n"
    contact_cell.text = contact_info
    set_cell_background(contact_cell, '87CEEB')  # Slightly darker gray background

    # Add content sections with titles
    titles = ['educational qualification','skillset and expertise', 'Previous work experience description','certifications','awards and achievements','Explanation of position of responsibilities', 'Explanation of projects', 'years of experience', 'previous job title']
    for title in titles:
        table = doc.add_table(rows=2, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(6.5)

        title_cell = table.cell(0, 0)
        # Apply bold to the title and set a light blue background
        run = title_cell.paragraphs[0].add_run(title.title())
        run.bold = True  # Convert title to title case
        set_cell_background(title_cell, 'ADD8E6')  # Dark gray background for title

        content_cell = table.cell(1, 0)
        # Retrieve content from json_data based on title
        content = json_data.get(title, 'Content not provided')
        print(title,content)
        # Special formatting for 'previous work experience description'
        if title == 'previous work experience description' and isinstance(content, list):
            for item in content:
            # Create a paragraph for each experience item
                p = content_cell.add_paragraph(style='ListBullet')
                # Split the item into subtitle and description
                job_title, _, description = item.partition(': ')
                # Add the job title as bold
                p.add_run(job_title + ': ').bold = True
                # Continue with the description
                p.add_run(description)
        elif title == 'skillset and expertise' and isinstance(content, list):
            # Join the skills with a comma and a space for the 'skillset and expertise' section
            content_cell.text = ', '.join(content)
        elif isinstance(content, list):
            # Add content as bullet points for list-type contents
            for item in content:
                content_cell.add_paragraph(item, style='ListBullet')
        elif content is None:  # Ensure content is not None
            content = []
        else:
            content_cell.text = content  # Directly add content if not a list
        
        print('/n \n')

        # Set the space after each table to zero
        set_space_after(table.rows[0].cells[0].paragraphs[0], 0)
        set_space_after(table.rows[1].cells[0].paragraphs[0], 0)


    # Save the document
    doc.save(filename)




# Function to set space after a paragraph to zero
def set_space_after2(paragraph, space):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Inches(space)

# Function for creating a document from a JSON-like structure
def create_doc_from_json2(json_data, filename):
    # Create a new Document
    doc = Document()
    
    # Set the page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Header with name and contact information
    doc.add_heading(json_data['name'], level=1)
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.add_run(f"Email: {json_data['gmail']} | ")
    contact_paragraph.add_run(f"Phone: {json_data['phone number']} | ")
    contact_paragraph.add_run(f"LinkedIn: {json_data['social media links'][0]}\n")
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_space_after2(contact_paragraph, 0.1)

    # Skillset and Expertise
    doc.add_heading('Skillset and Expertise', level=2)
    skills = ', '.join(json_data['skillset and expertise'])
    doc.add_paragraph(skills)
    
    # Certifications
    doc.add_heading('Certifications', level=2)
    for certification in json_data['certifications']:
        doc.add_paragraph(certification, style='ListBullet')
    
    # Projects
    doc.add_heading('Projects', level=2)
    for project in json_data['Explanation of projects']:
        doc.add_paragraph(project, style='ListBullet')
    
    # Positions of Responsibility
    doc.add_heading('Positions of Responsibility', level=2)
    for position in json_data['Explanation of position of responsibilities']:
        doc.add_paragraph(position, style='ListBullet')

    # Experience
    doc.add_heading('Experience', level=2)
    for experience in json_data['Previous work experience description']:
        doc.add_paragraph(experience, style='ListBullet')
    
    # Education
    doc.add_heading('Education', level=2)
    for education in json_data['educational qualification']:
        doc.add_paragraph(education, style='ListBullet')
    
    # Extracurriculars
    doc.add_heading('Extracurriculars', level=2)
    for activity in json_data['extracurriculars']:
        doc.add_paragraph(activity, style='ListBullet')
    
    # Awards and Achievements
    doc.add_heading('Awards and Achievements', level=2)
    for award in json_data['awards and achievements']:
        doc.add_paragraph(award, style='ListBullet')
    
    # Save the document
    doc.save(filename)


# Function to add a cell with a colored background
def set_cell_background(cell, fill):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Function to create a two-column table for layout
def create_two_column_table(doc, json_data):
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(5.5)  # Set the width of the left column
    table.columns[1].width = Cm(11.5) # Set the width of the right column

    # Left column content
    left_cell = table.cell(0, 0)
    set_cell_background(left_cell, "ADD8E6")  # Light blue background
    left_paragraph = left_cell.paragraphs[0]
    
    # Adding Name with a larger font to the left column
    run = left_paragraph.add_run(json_data['name'] + "\n")
    run.bold = True
    run.font.size = Pt(14)  # Set font size to 14pt or as desired
    
    # Adding Contact Information
    left_paragraph.add_run("Contact Information\n").bold = True
    left_paragraph.add_run(f"Email: {json_data['gmail']}\n")
    left_paragraph.add_run(f"Phone: {json_data['phone number']}\n")
    left_paragraph.add_run(f"LinkedIn: {json_data['social media links'][0]}\n")
    
    # Adding Educational Qualification with a bold title
    left_paragraph.add_run("\nEducational Qualification\n").bold = True
    for qualification in json_data['educational qualification']:
        left_paragraph.add_run(f"• {qualification}\n")

    # Adding Skills and Expertise with a bold title
    left_paragraph.add_run("\nSkillset and Expertise\n").bold = True
    # Join the skills with a comma and a space
    skills_text = ', '.join(json_data['skillset and expertise'])
    left_paragraph.add_run(f"• {skills_text}\n")
      
    # Adding previous job titles with a bold title
    left_paragraph.add_run("\nPrevious Job Roles\n").bold = True
    for skills in json_data['previous job title']:
        left_paragraph.add_run(f"• {skills}\n")

    # Adding years of experience with a bold title
    left_paragraph.add_run("\nYears of Experience\n").bold = True
    left_paragraph.add_run(f"• {json_data['years of experience']}\n")
    
    # Adding Certifications with a bold title
    left_paragraph.add_run("\nCertifications\n").bold = True
    for certification in json_data['certifications']:
        left_paragraph.add_run(f"• {certification}\n")

    # Right column content
    right_cell = table.cell(0, 1)
    right_paragraph = right_cell.paragraphs[0]
    
    # Adding Previous Work Experience Description with a bold title
    right_paragraph.add_run("Previous Work Experience Description\n").bold = True
    for experience in json_data['Previous work experience description']:
        right_paragraph.add_run(f"• {experience}\n")

    # Adding Projects with a bold title
    right_paragraph.add_run("\nProjects\n").bold = True
    for project in json_data['Explanation of projects']:
        right_paragraph.add_run(f"• {project}\n")

    # Adding Awards and Achievements with a bold title
    right_paragraph.add_run("\nAwards and Achievements\n").bold = True
    for award in json_data['awards and achievements']:
        right_paragraph.add_run(f"• {award}\n")

    # Adding Explanation of Position of Responsibilities with a bold title
    right_paragraph.add_run("\nExplanation of Position of Responsibilities\n").bold = True
    for responsibility in json_data['Explanation of position of responsibilities']:
        right_paragraph.add_run(f"• {responsibility}\n")
    
    # Adding Extracurricular Activities with a bold title
    right_paragraph.add_run("\nExtracurricular Activities\n").bold = True
    for activity in json_data['extracurriculars']:
        right_paragraph.add_run(f"• {activity}\n")

    return table

# Function for creating a document from a JSON-like structure
def create_doc_from_json3(json_data, filename):
    doc = Document()
    # Set the page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Create the two-column layout
    create_two_column_table(doc, json_data)

    # Save the document
    doc.save(filename)


# Function to download a docx file
def download_docx(doc, filename):
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes = docx_bytes.getvalue()
    st.download_button(label='Download Resume', data=docx_bytes, file_name=filename, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# Streamlit UI
st.title('Resume Parser and Builder')

uploaded_files = st.file_uploader("Upload your resumes here", type=["pdf", "doc", "docx", "txt"], accept_multiple_files=True)
# template_option = st.selectbox("Select Resume Template:", ("KGP format", "Simple format", "2 Column format"))

if uploaded_files:
    for uploaded_file in uploaded_files:
        with st.spinner(f'Parsing resume: {uploaded_file.name}...'):
            # Convert uploaded file to text
            resume_text = convert_file_to_text(uploaded_file)

            # Parse resume
            start_time = time.time()
            json_resume = await fetch_and_process(resume_text, systems)
            end_time = time.time()
            elapsed_time = end_time - start_time
            st.write(f"Execution time for {uploaded_file.name}: {elapsed_time:.2f} seconds")

            # Create and download resume for each template
            doc1_filename = f"Generated_Resume_KGP_{uploaded_file.name}.docx"
            create_doc_from_json_template1(json_resume, doc1_filename)
            download_docx(doc1_filename, doc1_filename)

            doc2_filename = f"Generated_Resume_Simple_{uploaded_file.name}.docx"
            create_doc_from_json_template2(json_resume, doc2_filename)
            download_docx(doc2_filename, doc2_filename)

            doc3_filename = f"Generated_Resume_2Column_{uploaded_file.name}.docx"
            create_doc_from_json_template3(json_resume, doc3_filename)
            download_docx(doc3_filename, doc3_filename)

